from fastapi import UploadFile
from pypdf import PdfReader
from docx import Document


def file_extractor(file: UploadFile):
    extension = file.filename.lower().split(".")[-1]
    if extension == "pdf":
        text = _pdf_extractor(file)
    elif extension == "docx":
        text = _docx_extractor(file)
    else:
        raise ValueError("Invalid file type!")
    return text


def _pdf_extractor(file: UploadFile):
    file.file.seek(0) #ensures that PdfReader or Document reads the file from the start
    reader = PdfReader(file.file)
    text = "\n".join(
        page.extract_text() or ""
        for page in reader.pages
    )
    return text

def _docx_extractor(file: UploadFile):
    file.file.seek(0)
    doc = Document(file.file)
    text = ""

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text.strip() + "\n"

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells
            )
            text += row_text + "\n"

    return text