from groq import Groq
from dotenv import load_dotenv
import os
from pypdf import PdfReader
from docx import Document
from pydantic import BaseModel
import json
from pathlib import Path
import time

load_dotenv()
API_KEY= os.getenv("GROQ_API_KEY")
MODEL = "llama-3.3-70b-versatile"

if not API_KEY:
    raise ValueError("api key not found")

client = Groq(api_key=API_KEY)

class Jd(BaseModel):
    role: str | None = None
    required_skills: list[str] = []
    preferred_skills: list[str] = []
    minimum_experience: str | None = None
    educational_requirements: list[str] = []
    responsibility: str | None = None

class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    tech_stack: list[str] = []
    duration: str | None = None
    description: str | None = None

class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone_no: str | None = None
    skills: list[str] = []
    experiences: list[Experience] = []
    projects: list[str] = []
    certifications: list[str] = []

class Result(BaseModel):
    score:float
    detail:dict

response_format={"type":"json_object"}

def get_jd(jd_text: str) -> Jd:
    schema = Jd.model_json_schema()

    system_prompt = f"""
You are an expert HR assistant specializing in analyzing job descriptions.

Your task is to extract all relevant information from the job description according to the provided schema.

Schema:
{json.dumps(schema, indent=2)}

Rules:
- Extract only explicitly stated information.
- Do not invent or assume any information.
- If a scalar field is missing, return null.
- If a list field is missing, return an empty list ([]).
- Return ONLY valid JSON.
- Do not include markdown or explanations.
"""

    messages = [
        {
            "role": "system",
            "content": system_prompt
        },
        {
            "role": "user",
            "content": jd_text
        }
    ]

    response = client.chat.completions.create(
        model=MODEL,
        messages=messages,
        response_format={"type": "json_object"}
    )

    data = json.loads(response.choices[0].message.content)

    return Jd(**data)

def file_extractor(filepath):
    if filepath.suffix.lower() == ".pdf":
        text=pdf_extractor(filepath)
    elif filepath.suffix.lower() == ".docx":
        text=docx_extractor(filepath)
    else:
        raise ValueError("invalid file type!! ")
    return text

def pdf_extractor(filepath):
    reader= PdfReader(filepath)
    text= "\n".join(
        page.extract_text() or ""
        for page in reader.pages
    )
    return text

def docx_extractor(filepath):
    doc = Document(filepath)

    text = ""

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text.strip() + "\n"

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            text += row_text + "\n"

    return text

def resume_parser(resume_text: str) -> Resume:
    schema = Resume.model_json_schema()

    system_prompt = f"""
You are an expert Resume Analyzer.

The user will provide a resume.

Extract all information according to the following schema.

Schema:
{json.dumps(schema, indent=2)}

Rules:
- Extract only information explicitly present in the resume.
- Do not invent or assume anything.
- If a scalar field is missing, return null.
- If a list field is missing, return an empty list ([]).
- Return ONLY valid JSON.
- Do not include markdown or explanations.
"""

    messages = [
        {
            "role": "system",
            "content": system_prompt
        },
        {
            "role": "user",
            "content": resume_text
        }
    ]

    response = client.chat.completions.create(
        model=MODEL,
        messages=messages,
        response_format={"type": "json_object"}
    )

    data = json.loads(response.choices[0].message.content)

    return Resume(**data)

def score_generator(jobd, resume):
    match_schema= Result.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {jobd.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """
    message={
        "role": "user",
        "content" : prompt
    }
    messages=[message]
    response_format={
        "type": "json_object"
    }
    response = client.chat.completions.create(model=MODEL, messages=messages, response_format=response_format)
    data = json.loads(response.choices[0].message.content)
    return Result(**data)

def main():
    all_results = []

    # -------------------------------
    # Read Job Description
    # -------------------------------
    jd_folder = Path("Job_description")

    raw_jd = None

    for file_path in jd_folder.iterdir():
        if file_path.suffix.lower() == ".txt":
            with open(file_path, "r", encoding="utf-8") as file:
                raw_jd = file.read()
        else:
            raw_jd = file_extractor(file_path)

        # Assuming there is only one JD
        break

    if raw_jd is None:
        print("No Job Description found.")
        return

    jd = get_jd(raw_jd)

    # -------------------------------
    # Process all resumes
    # -------------------------------
    resume_folder = Path("resumes")

    for file_path in resume_folder.iterdir():

        if not file_path.is_file():
            continue

        raw_resume = file_extractor(file_path)

        resume = resume_parser(raw_resume)

        time.sleep(5)

        result = score_generator(jd, resume)

        time.sleep(5)

        print(f"{resume.name}: {result.score}%")

        all_results.append({
            "name": resume.name or file_path.stem,
            "score": result.score,
            "detail": result.detail
        })

    # -------------------------------
    # Sort candidates
    # -------------------------------
    all_results.sort(
        key=lambda candidate: candidate["score"],
        reverse=True
    )

    top_1 = all_results[:1]
    worst_1 = all_results[-1:]

    # -------------------------------
    # Display Results
    # -------------------------------
    print("\n===== TOP CANDIDATE =====")

    for candidate in top_1:
        print(f"{candidate['name']} - {candidate['score']}%")
        print(candidate["detail"])

    print("\n===== LOWEST CANDIDATE =====")

    for candidate in worst_1:
        print(f"{candidate['name']} - {candidate['score']}%")
        print(candidate["detail"])


if __name__ == "__main__":
    main()